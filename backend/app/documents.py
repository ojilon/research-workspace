"""
Document extraction → structured block tree.

Each block is a selectable unit (paragraph, heading, list item, table cell text).
This powers:
  - click-to-select a whole paragraph (no drag-select required)
  - future drag-and-drop into a summary pane
  - future search / autocorrect over blocks

Supported:
  - PDF   via pypdf (text layer; scanned PDFs need OCR later)
  - DOCX  via python-docx (paragraphs + tables)
  - DOC   not supported natively (save as DOCX in Word)
  - plain text / markdown / code → one block per non-empty line or whole file
"""

from __future__ import annotations

import uuid
from pathlib import Path
from typing import Any


def _uid(prefix: str = "b") -> str:
    return f"{prefix}-{uuid.uuid4().hex[:10]}"


def file_kind(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return "pdf"
    if ext == ".docx":
        return "docx"
    if ext == ".doc":
        return "doc"
    if ext in {".md", ".markdown", ".txt"}:
        return "text"
    if ext in {
        ".py",
        ".ts",
        ".tsx",
        ".js",
        ".jsx",
        ".json",
        ".css",
        ".html",
        ".rs",
        ".go",
        ".java",
        ".c",
        ".cpp",
        ".h",
        ".sh",
        ".yml",
        ".yaml",
        ".toml",
        ".sql",
    }:
        return "code"
    if ext in {".url", ".webloc"}:
        return "link"
    return "binary"


def extract_document(path: Path) -> dict[str, Any]:
    """
    Return a document payload:
      {
        "path": str,
        "kind": "pdf"|"docx"|...,
        "title": str,
        "blocks": [ { id, type, text, meta? }, ... ],
        "warning": optional str
      }
    """
    kind = file_kind(path)
    title = path.stem

    if kind == "pdf":
        return _extract_pdf(path, title)
    if kind == "docx":
        return _extract_docx(path, title)
    if kind == "doc":
        return {
            "path": str(path),
            "kind": "doc",
            "title": title,
            "blocks": [],
            "warning": "Legacy .doc is not supported. Open in Word and save as .docx.",
        }
    if kind in {"text", "code"}:
        return _extract_text_file(path, title, kind)

    return {
        "path": str(path),
        "kind": kind,
        "title": title,
        "blocks": [],
        "warning": f"No text extractor for this file type ({path.suffix}).",
    }


def _extract_pdf(path: Path, title: str) -> dict[str, Any]:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    blocks: list[dict[str, Any]] = []
    for page_index, page in enumerate(reader.pages):
        text = (page.extract_text() or "").strip()
        if not text:
            blocks.append(
                {
                    "id": _uid(),
                    "type": "paragraph",
                    "text": "",
                    "meta": {"page": page_index + 1, "empty": True},
                }
            )
            continue
        # Split into paragraphs on blank lines for selectable units
        parts = [p.strip() for p in text.split("\n\n") if p.strip()]
        if not parts:
            parts = [text]
        for part in parts:
            # Also split very long single newlines into softer paragraphs
            for line_group in _chunk_lines(part):
                blocks.append(
                    {
                        "id": _uid(),
                        "type": "paragraph",
                        "text": line_group,
                        "meta": {"page": page_index + 1},
                    }
                )

    warning = None
    if not any(b["text"] for b in blocks):
        warning = (
            "No extractable text (scanned PDF?). OCR can be added later; "
            "you can still view the PDF visually."
        )

    return {
        "path": str(path),
        "kind": "pdf",
        "title": title,
        "blocks": blocks,
        "page_count": len(reader.pages),
        "warning": warning,
    }


def _chunk_lines(text: str, max_lines: int = 8) -> list[str]:
    lines = text.splitlines()
    if len(lines) <= max_lines:
        return [text]
    out: list[str] = []
    for i in range(0, len(lines), max_lines):
        out.append("\n".join(lines[i : i + max_lines]).strip())
    return [x for x in out if x]


def _extract_docx(path: Path, title: str) -> dict[str, Any]:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE

    doc = Document(str(path))
    blocks: list[dict[str, Any]] = []

    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        style_name = ""
        try:
            style_name = para.style.name if para.style else ""
        except Exception:
            style_name = ""

        block_type = "paragraph"
        level = None
        if style_name.startswith("Heading"):
            block_type = "heading"
            try:
                level = int(style_name.replace("Heading", "").strip() or "1")
            except ValueError:
                level = 1
        elif "List" in style_name:
            block_type = "list_item"

        runs_meta = []
        for run in para.runs:
            runs_meta.append(
                {
                    "text": run.text,
                    "bold": bool(run.bold),
                    "italic": bool(run.italic),
                    "underline": bool(run.underline),
                }
            )

        blocks.append(
            {
                "id": _uid(),
                "type": block_type,
                "text": text,
                "meta": {
                    "style": style_name,
                    "level": level,
                    "runs": runs_meta,
                },
            }
        )

    # Tables → one block per cell with text
    for t_index, table in enumerate(doc.tables):
        for r_index, row in enumerate(table.rows):
            for c_index, cell in enumerate(row.cells):
                cell_text = (cell.text or "").strip()
                if not cell_text:
                    continue
                blocks.append(
                    {
                        "id": _uid(),
                        "type": "table_cell",
                        "text": cell_text,
                        "meta": {
                            "table": t_index,
                            "row": r_index,
                            "col": c_index,
                        },
                    }
                )

    return {
        "path": str(path),
        "kind": "docx",
        "title": title,
        "blocks": blocks,
        "warning": None if blocks else "DOCX had no extractable paragraphs.",
    }


def _extract_text_file(path: Path, title: str, kind: str) -> dict[str, Any]:
    try:
        raw = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        raw = path.read_text(encoding="utf-8", errors="replace")

    if kind == "code":
        lines = raw.splitlines()
        blocks = [
            {
                "id": _uid(),
                "type": "code_line",
                "text": line,
                "meta": {"line": i + 1},
            }
            for i, line in enumerate(lines)
        ]
        return {
            "path": str(path),
            "kind": "code",
            "title": title,
            "blocks": blocks,
            "raw": raw,
            "language": path.suffix.lstrip("."),
        }

    # Markdown / plain text → paragraph blocks
    parts = [p.strip() for p in raw.split("\n\n") if p.strip()]
    if not parts and raw.strip():
        parts = [raw.strip()]
    blocks = [
        {"id": _uid(), "type": "paragraph", "text": p, "meta": {}}
        for p in parts
    ]
    return {
        "path": str(path),
        "kind": "text",
        "title": title,
        "blocks": blocks,
        "raw": raw,
    }
