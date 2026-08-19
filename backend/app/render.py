"""
Lightweight HTML rendering helpers for documents.

Python extracts structure; the frontend displays the HTML.
Not a full Word clone — enough to read with headings, lists, tables.
"""

from __future__ import annotations

import html
from pathlib import Path


def docx_to_html(path: Path) -> str:
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(str(path))
    parts: list[str] = [
        "<!DOCTYPE html><html><head><meta charset='utf-8'>",
        "<style>",
        "body{font-family:Georgia,'Times New Roman',serif;line-height:1.55;"
        "max-width:42rem;margin:1.5rem auto;padding:0 1.25rem;color:#1a1a1a;}",
        "h1{font-size:1.6rem;margin:1.4rem 0 .6rem}",
        "h2{font-size:1.35rem;margin:1.2rem 0 .5rem}",
        "h3{font-size:1.15rem;margin:1rem 0 .4rem}",
        "p{margin:.55rem 0}",
        "ul,ol{margin:.5rem 0 .5rem 1.25rem}",
        "table{border-collapse:collapse;width:100%;margin:1rem 0;font-size:.95rem}",
        "td,th{border:1px solid #ccc;padding:.4rem .55rem;vertical-align:top}",
        "th{background:#f3f3f3}",
        "@media (prefers-color-scheme:dark){"
        "body{background:#1e1e1e;color:#e8e8e8}"
        "td,th{border-color:#444}th{background:#2a2a2a}}",
        "</style></head><body>",
    ]

    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        style = ""
        try:
            style = para.style.name if para.style else ""
        except Exception:
            style = ""

        safe = _runs_to_html(para)
        if style.startswith("Heading"):
            try:
                level = int("".join(c for c in style if c.isdigit()) or "1")
            except ValueError:
                level = 1
            level = min(max(level, 1), 4)
            parts.append(f"<h{level}>{safe}</h{level}>")
        elif "List" in style:
            parts.append(f"<ul><li>{safe}</li></ul>")
        else:
            parts.append(f"<p>{safe}</p>")

    for table in doc.tables:
        parts.append("<table>")
        for i, row in enumerate(table.rows):
            parts.append("<tr>")
            tag = "th" if i == 0 else "td"
            for cell in row.cells:
                cell_html = html.escape((cell.text or "").strip()).replace("\n", "<br>")
                parts.append(f"<{tag}>{cell_html}</{tag}>")
            parts.append("</tr>")
        parts.append("</table>")

    parts.append("</body></html>")
    return "".join(parts)


def _runs_to_html(para) -> str:
    """Preserve basic bold/italic from runs when present."""
    if not para.runs:
        return html.escape(para.text or "")
    chunks: list[str] = []
    for run in para.runs:
        t = html.escape(run.text or "")
        if not t:
            continue
        if run.bold:
            t = f"<strong>{t}</strong>"
        if run.italic:
            t = f"<em>{t}</em>"
        if run.underline:
            t = f"<u>{t}</u>"
        chunks.append(t)
    return "".join(chunks) if chunks else html.escape(para.text or "")
