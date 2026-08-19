"""
Open Windows Explorer or a native Save As dialog.
"""

from __future__ import annotations

import os
import subprocess
from pathlib import Path


def open_in_explorer(path: Path, select: bool = False) -> None:
    path = path.resolve()
    if os.name == "nt":
        if select and path.is_file():
            subprocess.Popen(["explorer", f"/select,{path}"])
        elif path.is_dir():
            subprocess.Popen(["explorer", str(path)])
        else:
            target = path if path.is_dir() else path.parent
            target.mkdir(parents=True, exist_ok=True)
            subprocess.Popen(["explorer", str(target)])
        return

    opener = "open" if os.name == "posix" and Path("/usr/bin/open").exists() else "xdg-open"
    subprocess.Popen([opener, str(path if path.is_dir() else path.parent)])


def save_as_dialog(
    initial_dir: Path,
    default_name: str = "Untitled.md",
    title: str = "Create document",
) -> Path | None:
    """
    Show a native Save As dialog (Windows via tkinter).

    Returns the chosen path, or None if the user cancelled.
    Creates an empty file if it does not exist yet.
    """
    initial_dir = initial_dir.resolve()
    initial_dir.mkdir(parents=True, exist_ok=True)

    try:
        import tkinter as tk
        from tkinter import filedialog
    except ImportError:
        # No GUI toolkit – fall back to a default path
        target = initial_dir / default_name
        if not target.exists():
            target.write_text("", encoding="utf-8")
        return target

    root = tk.Tk()
    root.withdraw()
    root.attributes("-topmost", True)

    filetypes = [
        ("Markdown", "*.md"),
        ("Text", "*.txt"),
        ("Word document", "*.docx"),
        ("All files", "*.*"),
    ]

    chosen = filedialog.asksaveasfilename(
        parent=root,
        title=title,
        initialdir=str(initial_dir),
        initialfile=default_name,
        defaultextension=".md",
        filetypes=filetypes,
    )
    root.destroy()

    if not chosen:
        return None

    path = Path(chosen)
    if not path.exists():
        # Seed empty content by extension
        if path.suffix.lower() == ".docx":
            _create_empty_docx(path)
        else:
            path.parent.mkdir(parents=True, exist_ok=True)
            path.write_text("# " + path.stem + "\n\n", encoding="utf-8")
    return path.resolve()


def _create_empty_docx(path: Path) -> None:
    try:
        from docx import Document

        doc = Document()
        doc.add_heading(path.stem, level=1)
        doc.add_paragraph("")
        path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(path))
    except Exception:
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_bytes(b"")
